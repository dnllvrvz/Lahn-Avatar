import os
import subprocess
import datetime, requests
from docx import Document
from urllib.parse import urlparse, parse_qs
from pathlib import Path
import hashlib
import re
from dotenv import load_dotenv

from youtube_transcript_api import YouTubeTranscriptApi
from llama_index.core.schema import Document as LlamaDocument

from llama_index.core import StorageContext, load_index_from_storage, Settings
from llama_index.core.readers import SimpleDirectoryReader
from llama_index.core.indices.vector_store import VectorStoreIndex
from llama_index.readers.web import SimpleWebPageReader


from openai import OpenAI


from llama_index.embeddings.huggingface import HuggingFaceEmbedding

from .gwdg_llm import GWDGChatLLM


load_dotenv()


# === CONFIG ===
DRIVE_FOLDER_ID = "1vT4UTYHeFxS5Vy2u_OfQyQ6cQ-cP5Ywd"
API_KEY = os.getenv("GWDG_API_KEY") 
API_BASE = os.getenv("GWDG_API_BASE")

AZURE_VERSION = os.getenv("AZURE_API_VERSION")
AZURE_KEY = os.getenv("AZURE_CHAT_KEY")
AZURE_BASE = os.getenv("AZURE_API_BASE")

# base_dir = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = "./data" #os.path.join(base_dir, "/data")

# print('Base dir: ', base_dir, 'Data dir: ', DATA_DIR)
LOG_DIR = "./chat_logs" #os.path.join(base_dir, "/chat_logs")
STORAGE_DIR = "./lahn_index"#os.path.join(base_dir, "/lahn_index")


def download_drive_folder(folder_id, output_dir="./data"):
    print('Running download_drive_folder function...')
    os.makedirs(output_dir, exist_ok=True)
    cmd = f"gdown --folder https://drive.google.com/drive/folders/{folder_id} -O {output_dir}"
    subprocess.run(cmd, shell=True)


def fetch_system_prompt_from_gdoc():
    print(' Updating system prompt...')
    url = "https://docs.google.com/document/d/1NYOOy8KkaLDBwvHvEVg1hVDY5yvHeLACUpCEkJVM8Kw/export?format=txt"
    response = requests.get(url)
    response.raise_for_status()
    prompt = response.text.strip()
    # prompt = prompt[:prompt.find('General Internal Impressions')]

    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'system_prompt.txt')
    with open(file_path, 'w') as f:
        f.write(prompt)
    print(' Done.')


def convert_docx_to_txt_and_cleanup(folder_path):
    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith('.docx') or '.' not in file:
                try:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    txt_filename = os.path.splitext(file)[0] + '.txt'
                    txt_path = os.path.join(root, txt_filename)
                    with open(txt_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    os.remove(file_path)
                    print(f"✅ Converted and deleted: {file_path}")
                except Exception as e:
                    print(f"❌ Failed to convert {file_path}: {e}")


def fetch_youtube_transcript(url, languages=["de"]):
    print(f"🔗 Fetching: {url}")
    try:
        parsed = urlparse(url)

        # Handle standard and short YouTube URL formats
        if "youtube.com" in parsed.netloc:
            video_id = parse_qs(parsed.query).get("v", [None])[0]
        elif "youtu.be" in parsed.netloc:
            video_id = parsed.path.lstrip("/")
        else:
            raise ValueError("Unrecognized YouTube URL format.")

        if not video_id or len(video_id) != 11:
            raise ValueError("Invalid YouTube video ID.")

        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=languages)
        full_text = " ".join([entry["text"] for entry in transcript])
        return LlamaDocument(text=full_text, metadata={"source": url})

    except Exception as e:
        print(f"❌ Failed to fetch {url}: {e}")
        return None


def sanitize_filename(url):
    domain = urlparse(url).netloc
    hashed = hashlib.md5(url.encode()).hexdigest()[:8]
    return f"{domain.replace('.', '_')}_{hashed}.txt"



def select_model():
    print("Choose a model:")
    print("1. Mistral")
    print("2. SauerKrautLM (Llama finetuned on German data)")
    choice = input("Enter 1 or 2: ")
    return "llama-3.1-sauerkrautlm-70b-instruct" if choice == "2" else "mistral-large-instruct"


def get_llm(mode='openai',model_name=None, system_prompt=None):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'system_prompt.txt')
    if system_prompt == None:
        system_prompt = open(file_path, 'r').read()

    # if model_name != None:
    if mode == 'openai':

        llm =  OpenAI(
            # model=model_name,        # your HRZ model name
            # temperature=0.5,
            # system_prompt=system_prompt,
            # context_window=128000,

            # point at your custom endpoint:
            api_key=API_KEY,             # e.g. 'sk-…'
            base_url=API_BASE,           # "https://llm.hrz.uni-giessen.de/api/"
            # api_type="open_ai",          # use the “open_ai” protocol
            # api_version=None,            # leave None unless your server needs a version
        )


    else:

        llm = GWDGChatLLM(
                model=model_name,
                api_base=API_BASE,
                api_key=API_KEY,
                temperature=0.5,
                system_prompt=system_prompt
            )

    # print('LLM details: ', llm.model_dump())

    # Settings.llm = llm

    return llm, system_prompt


def create_session_log():
    os.makedirs(LOG_DIR, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    return open(os.path.join(LOG_DIR, f"session_{timestamp}.txt"), "w")


#wrt Text Index

# pip install deep_translator langdetect rank_bm25 nltk unicodedata
import re, unicodedata, pickle
from langdetect import detect
from deep_translator import GoogleTranslator   # or DeeplTranslator
from nltk.tokenize import word_tokenize, sent_tokenize
from rank_bm25 import BM25Okapi
import nltk

def ensure_punkt_tab():
    try:
        # this will raise LookupError if not found
        nltk.data.find('tokenizers/punkt_tab')
    except LookupError:
        print("punkt_tab not found—downloading…")
        nltk.download('punkt_tab')
    else:
        print("punkt_tab already available.")


CHUNK_SIZE, OVERLAP = 200, 30
# ------------------------------------------------------------------
# 0. normalisers / helpers
# ------------------------------------------------------------------
def normalise(txt:str) -> str:
    txt = re.sub(r"\s+", " ", txt.lower().strip())
    return unicodedata.normalize("NFKD", txt)

def tokenize(text:str, lang:str) -> list[str]:
    lang_flag = "german" if lang=="de" else "english"
    return word_tokenize(text, language=lang_flag)

# ------------------------------------------------------------------
# 1. translation helper with cache
# ------------------------------------------------------------------
_trans_cache: dict[tuple[str,str], str] = {}   # (text, target_lang) → translated

def translate(text:str, target_lang:str) -> str:
    """
    Translate 'text' to target language ('de' or 'en') using deep_translator.
    Caches results to avoid hitting rate limits.
    """
    key = (text, target_lang)
    if key in _trans_cache:
        return _trans_cache[key]

    translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
    _trans_cache[key] = translated
    return translated


def prepare_text_index(RAW_TEXT):
    ensure_punkt_tab()
    sentences = sent_tokenize(normalise(RAW_TEXT), language="german")
    chunks, cur, wc = [], [], 0
    for sent in sentences:
        words = tokenize(sent, "de")      # punkt model is DE but works for EN too
        cur.extend(words); wc += len(words)
        if wc >= CHUNK_SIZE:
            chunks.append(" ".join(cur))
            cur, wc = cur[-OVERLAP:], len(cur)
    if cur: chunks.append(" ".join(cur))

    print(f"➡  Raw chunks: {len(chunks)}")
    # ------------------------------------------------------------------
    # 2. Build your BM-25 index (assumes you already have 'chunks')
    # ------------------------------------------------------------------
    #   chunks = ["..."]    # list of cleaned 200-word blocks
    token_lists = [tokenize(c, detect(c)) for c in chunks]
    bm25 = BM25Okapi(token_lists)
    return bm25, chunks

# ------------------------------------------------------------------
# 3. Dual-language search
# ------------------------------------------------------------------
def search_text_index(bm25, chunks, query:str, k_each:int=5):
    keyword_list = query.split(', ')
    query = ' '.join(keyword_list[:3])
    trans_q = ' '.join(keyword_list[3:])

    lang_orig  = 'en' #"de" if detect(query) == "de" else "en"
    lang_trans = 'de' #"en" if lang_orig == "de" else "de"

    # --- original language pass -----------------------------------
    q_tokens_o = tokenize(normalise(query), lang_orig)
    print('Query recieved by Text Index searcher: ', ', '.join(keyword_list))
    print('Keywords group 1: ', query)
    print('Keywords group 2: ', trans_q)

    print('Group 1 tokens to search with BM25: ', q_tokens_o)
    scores_o   = bm25.get_scores(q_tokens_o)
    top_o      = scores_o.argsort()[-k_each:][::-1]

    # --- translated pass ------------------------------------------
    # trans_q    = translate(query, lang_trans)
    print('Translated query: ',trans_q)
    q_tokens_t = tokenize(normalise(trans_q), lang_trans)
    print('Group 2 tokens to search with BM25: ', q_tokens_t)
    scores_t   = bm25.get_scores(q_tokens_t)
    top_t      = scores_t.argsort()[-k_each:][::-1]

    # --- merge, preferring best score if overlap ------------------
    seen, results = {}, []
    for idx in top_o:
        seen[idx] = ("orig", float(scores_o[idx]))
    for idx in top_t:
        if idx in seen:
            # keep the better score
            seen[idx] = ("orig+trans", max(seen[idx][1], float(scores_t[idx])))
        else:
            seen[idx] = ("trans", float(scores_t[idx]))

    # sort by score descending and trim to k_each*2
    merged = sorted(seen.items(), key=lambda kv: kv[1][1], reverse=True)[: 2*k_each]
    for idx, (tag, score) in merged[:6]:
        # results.append((tag, score, chunks[idx]))
        results.append(chunks[idx])
    return results

# ------------------------------------------------------------------
# 4. demo
# ------------------------------------------------------------------
# for tag, score, snippet in search_dual("Tell me about fish and their sizes"):
#     print(f"[{tag}] {score:6.2f}  {snippet}")








from langdetect import detect
from deep_translator import GoogleTranslator
import spacy

# Load multilingual models (download with: python -m spacy download en_core_web_sm && python -m spacy download de_core_news_sm)
nlp_en = spacy.load("en_core_web_sm")
nlp_de = spacy.load("de_core_news_sm")




def translate_keywords_batch(keywords, source_lang="auto", target_lang="de"):
    # Join all keywords into a comma-separated string
    joined = ", ".join(keywords)
    translated_text = GoogleTranslator(source=source_lang, target=target_lang).translate(joined)
    # Split back into individual words
    translated_keywords = [w.strip() for w in translated_text.split(",")]
    return translated_keywords


def extract_keywords(text):
    # Detect language
    lang = detect(text)
    if lang.startswith('de'):
        nlp = nlp_de
        target_lang = 'en' #Assumes the language is English if it's nt German
    else:
        nlp = nlp_en
        target_lang = 'de'

    doc = nlp(text)
    # Extract nouns and proper nouns as keywords
    keywords = [token.text for token in doc if token.pos_ in ("NOUN", "PROPN") and not token.is_stop]
    keywords = list(set(keywords))  # remove duplicates

    translated_keywords = translate_keywords_batch(keywords, target_lang=target_lang)

    if target_lang == 'en':
        all_keywords = translated_keywords + keywords
    else:
        all_keywords = keywords + translated_keywords

    keyword_list = ', '.join(all_keywords)

    return keyword_list







def RAG(query):
    keywords_for_text_based_retrieval = extract_keywords(query)
    print('Keywords for text-based retrieval: ', keywords_for_text_based_retrieval)

    with ThreadPoolExecutor(max_workers=2) as executor:
        thread_0 = executor.submit(fetch_vector_index_context, query)
        thread_1 = executor.submit(text_index_query_engine, text_index, chunks, keywords_for_text_based_retrieval)

        wait([thread_0,thread_1])

    # context_from_text_index = thread_0
    context_from_vector_index =  thread_0.result().response
    context_from_text_index = thread_1.result()

    total_context = '\nContext from text-based retrieval: \n' +context_from_text_index + '\n------------\nContext from vector-based retrieval: \n' + context_from_vector_index

    total_context = 'Here is relevant information about the Lahn (Sometimes the text-retrieval has relevant information that the vector-retrieval doesn\'t, or vice versa. Look through each comprehensively, to extract the information you need. Even if the Vector-retrieval says there\'s no information available, still scrutinize the Text-retrieval results to fetch relevant info: ' + total_context

    print('RAG result: ', total_context)
    return total_context


def build_index():
    # cmd = f"shopt -s extglob; rm -rf data/!(uploaded_experiences)"
    # subprocess.run(cmd, shell=True) #

    subprocess.run(
        "find data -mindepth 1 -not -name 'uploaded_experiences' -exec rm -rf {} +",
        shell=True,
        check=True
    )


    print('Just cleared data/ . Contents: ', os.listdir('data'))


    print('Refreshing from Google Drive...')
    download_drive_folder(DRIVE_FOLDER_ID, DATA_DIR)
    convert_docx_to_txt_and_cleanup(DATA_DIR)

    print('Creating Context store from data sources...')

    if len(os.listdir(DATA_DIR))>0:
        documents = SimpleDirectoryReader(DATA_DIR, recursive=True).load_data()
        print(f"{len(documents)} documents loaded from {DATA_DIR}")
        # for i, doc in enumerate(documents):
        #     print(f"\n--- Document {i+1} ---")
        #     print("File:", doc.metadata.get('file_path', 'Unknown'))
        #     print("Content preview:", doc.text[:300], "...\n")


    links_path = Path(DATA_DIR) / "General_News/Online News (Links).txt"
    if links_path.exists():
        with open(links_path, "r") as f:
            urls = [line.strip() for line in f if line.strip()]

        web_reader = SimpleWebPageReader()
        for url in urls:
            try:
                print(f"🔗 Fetching: {url}")
                if "youtube.com" in url or "youtu.be" in url:
                    doc = fetch_youtube_transcript(url)
                else:
                    docs = web_reader.load_data([url])
                    full_text = "\n\n".join(doc.text for doc in docs)
                    doc = LlamaDocument(text=full_text, metadata={"source": url})

                if doc:
                    filename = sanitize_filename(url)
                    filepath = Path(DATA_DIR) / "General_News/scraped_texts" / filename
                    filepath.parent.mkdir(parents=True, exist_ok=True)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(doc.text)
                    print(f"✅ Saved to {filepath}")
            except Exception as e:
                print(f"❌ Failed to fetch {url}: {e}")

    scraped_documents_path = Path(DATA_DIR) / "General_News/scraped_texts"
    if  scraped_documents_path.exists() and len(os.listdir(str(scraped_documents_path)))>0:
        scraped_documents = SimpleDirectoryReader(str(Path(DATA_DIR) / "General_News/scraped_texts")).load_data()
        documents += scraped_documents

    experiences_folder_path = Path(DATA_DIR) / "uploaded_experiences/text"
    # experiences_folder_is_empty = not os.listdir(str(Path(DATA_DIR) / "uploaded_experiences/text"))
    if experiences_folder_path.exists() and len(os.listdir(str(experiences_folder_path)))>0:
        new_uploads = SimpleDirectoryReader(str(Path(DATA_DIR) / "uploaded_experiences"), recursive=True).load_data()
        documents += new_uploads
        # user_experiences = SimpleDirectoryReader(str(Path(DATA_DIR) / "uploaded_experiences/text")).load_data()
        # documents += user_experiences

    vector_index = VectorStoreIndex.from_documents(documents)
    vector_index.storage_context.persist(persist_dir=STORAGE_DIR)


    context = '\n'.join([doc.text for doc in documents])
    text_index, chunks = prepare_text_index(context)

    pickle.dump(text_index, open(STORAGE_DIR+'/text_index.pkl','wb'))
    pickle.dump(chunks, open(STORAGE_DIR+'/chunks.pkl','wb'))
    # index = text_index


    print('Done')

    return vector_index, text_index, chunks



def build_or_load_index(refresh=False):
    Settings.embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

    index_ready = (
        os.path.exists(STORAGE_DIR)

        #Vector index
        and os.path.exists(os.path.join(STORAGE_DIR, "docstore.json"))
        and os.path.exists(os.path.join(STORAGE_DIR, "index_store.json"))

        #Text index
        and os.path.exists(os.path.join(STORAGE_DIR, "text_index.pkl"))
        and os.path.exists(os.path.join(STORAGE_DIR, "chunks.pkl"))
    )

    if index_ready and not refresh:
        print('Loading index from storage...')
        storage_context = StorageContext.from_defaults(persist_dir=STORAGE_DIR)
        vector_index = load_index_from_storage(storage_context)

        text_index = pickle.load(open(STORAGE_DIR+'/text_index.pkl','rb'))
        chunks = pickle.load(open(STORAGE_DIR+'/chunks.pkl','rb'))
        return vector_index, text_index, chunks

    #Index needs to be built and loaded
    vector_index, text_index, chunks = build_index()
    
    return vector_index, text_index, chunks







# #For caching translations relevant to extracting context from Text-based Index:

# import hashlib
# from deep_translator import GoogleTranslator

# # In-memory translation cache
# _trans_cache = {} #Save to file, for persistence.

# def _cache_key(word: str, target_lang: str) -> str:
#     """
#     Generate a hash key for each word-language pair.
#     """
#     key_str = f"{word.lower().strip()}::{target_lang}"
#     return hashlib.md5(key_str.encode("utf-8")).hexdigest()

# def translate_batch(words, target_lang: str):
#     """
#     Translates a list of words to target_lang, using per-word caching.
#     Only uncached words trigger API calls.
#     """
#     results = []
#     uncached_words = []
#     uncached_indices = []

#     # 1️⃣ Check cache for each word
#     for i, word in enumerate(words):
#         key = _cache_key(word, target_lang)
#         if key in _trans_cache:
#             results.append(_trans_cache[key])
#         else:
#             results.append(None)
#             uncached_words.append(word)
#             uncached_indices.append(i)

#     # 2️⃣ Translate uncached words in a batch (if API allows)
#     if uncached_words:
#         # GoogleTranslator can take a single string — we join by newline for batch translation
#         text_to_translate = "\n".join(uncached_words)
#         translated_text = GoogleTranslator(source="auto", target=target_lang).translate(text_to_translate)
        
#         # deep_translator may return one big string, so split it back
#         translated_list = [t.strip() for t in translated_text.split("\n") if t.strip()]

#         # Safety: align lengths
#         if len(translated_list) != len(uncached_words):
#             raise ValueError("Batch translation mismatch between input and output sizes.")

#         # 3️⃣ Store each translation in cache and fill results
#         for idx, word, translated in zip(uncached_indices, uncached_words, translated_list):
#             key = _cache_key(word, target_lang)
#             _trans_cache[key] = translated
#             results[idx] = translated

#             # optional reverse cache
#             reverse_key = _cache_key(translated, "en" if target_lang != "en" else "auto")
#             _trans_cache[reverse_key] = word

#     return results
