# ---------- standard library ----------
import datetime
import hashlib
import io
import json
import os
import pickle
import re
import subprocess
import time
import unicodedata
import wave
from concurrent.futures import ThreadPoolExecutor, wait
from pathlib import Path
from typing import Any, List
from urllib.parse import parse_qs, urlparse

import langid
import nltk
import pandas as pd

# ---------- third-party ----------
import requests
import spacy
import torch
import torchaudio
from deep_translator import GoogleTranslator
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from langdetect import detect

# ---------- llama-index ----------
from llama_index.core import Settings, StorageContext, load_index_from_storage
from llama_index.core.indices.vector_store import VectorStoreIndex
from llama_index.core.readers import SimpleDirectoryReader
from llama_index.core.schema import Document as LlamaDocument
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.experimental.query_engine import PandasQueryEngine
from llama_index.readers.web import SimpleWebPageReader
from nltk.tokenize import sent_tokenize, word_tokenize
from openai import OpenAI
from rank_bm25 import BM25Okapi
from youtube_transcript_api import YouTubeTranscriptApi

# --- LLAMA-INDEX GLOBAL SETTINGS ---
# Load the embedding model once and set it globally
Settings.embed_model = HuggingFaceEmbedding(
    model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
)
# ---

# === CONFIG ===


# base_dir = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = "./data"  # os.path.join(base_dir, "/data")

# print('Base dir: ', base_dir, 'Data dir: ', DATA_DIR)
LOG_DIR = "./chat_logs"  # os.path.join(base_dir, "/chat_logs")
STORAGE_DIR = "./lahn_index"  # os.path.join(base_dir, "/lahn_index")


def download_drive_folder(folder_id, output_dir="./data"):
    print("Running download_drive_folder function using Google Drive API...")
    os.makedirs(output_dir, exist_ok=True)
    drive_service = _build_drive_service()
    if not drive_service:
        print("Failed to initialize Google Drive service. Aborting download.")
        return

    def _download_recursive(current_folder_id, current_output_path):
        page_token = None
        while True:
            # List files and folders in the current Google Drive folder
            response = (
                drive_service.files()
                .list(
                    q=f"'{current_folder_id}' in parents and trashed=false",
                    fields="nextPageToken, files(id, name, mimeType)",
                    pageToken=page_token,
                )
                .execute()
            )

            for item in response.get("files", []):
                file_id = item["id"]
                file_name = item["name"]
                mime_type = item["mimeType"]
                output_file_path = os.path.join(current_output_path, file_name)

                if mime_type == GDRIVE_FOLDER_MIMETYPE:
                    # If it's a folder, create local directory and recurse
                    os.makedirs(output_file_path, exist_ok=True)
                    print(f"Created local folder: {output_file_path}")
                    _download_recursive(file_id, output_file_path)
                else:
                    # It's a file, download it
                    print(
                        f"Downloading file: {file_name} (ID: {file_id}) to {current_output_path}"
                    )
                    try:
                        request = None
                        if mime_type == GDOCS_MIMETYPE:
                            # Google Docs export as plain text
                            request = drive_service.files().export_media(
                                fileId=file_id, mimeType="text/plain"
                            )
                            output_file_path += ".txt"  # Ensure it has a .txt extension
                        elif mime_type == GSHEETS_MIMETYPE:
                            # Google Sheets export as CSV
                            request = drive_service.files().export_media(
                                fileId=file_id, mimeType="text/csv"
                            )
                            output_file_path += ".csv"  # Ensure it has a .csv extension
                        elif mime_type == GPTS_MIMETYPE:
                            # Google Slides export as PDF
                            request = drive_service.files().export_media(
                                fileId=file_id, mimeType="application/pdf"
                            )
                            output_file_path += ".pdf"  # Ensure it has a .pdf extension
                        else:
                            # Regular file download
                            request = drive_service.files().get_media(fileId=file_id)

                        if request:
                            with open(output_file_path, "wb") as fh:
                                downloader = MediaIoBaseDownload(fh, request)
                                done = False
                                while done is False:
                                    status, done = downloader.next_chunk()
                                    # print(f"Download progress for {file_name}: {int(status.progress() * 100)}%")
                            print(f"Successfully downloaded: {output_file_path}")
                        else:
                            print(
                                f"Skipping unknown MIME type for file: {file_name} ({mime_type})"
                            )
                    except Exception as e:
                        print(f"Error downloading {file_name}: {e}")

            page_token = response.get("nextPageToken", None)
            if not page_token:
                break

    _download_recursive(folder_id, output_dir)
    print("Google Drive API download complete.")


GDOCS_MIMETYPE = "application/vnd.google-apps.document"
GSHEETS_MIMETYPE = "application/vnd.google-apps.spreadsheet"
GPTS_MIMETYPE = "application/vnd.google-apps.presentation"
GDRIVE_FOLDER_MIMETYPE = "application/vnd.google-apps.folder"


def _build_drive_service():
    """Builds and returns a Google Drive API service object using service account credentials."""
    SERVICE_ACCOUNT_FILE = "utils/.gdrive_service_account"
    SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

    try:
        creds = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES
        )
        return build("drive", "v3", credentials=creds)
    except Exception as e:
        print(f"Error building Drive service: {e}")
        return None


# ="https://docs.google.com/document/d/1NYOOy8KkaLDBwvHvEVg1hVDY5yvHeLACUpCEkJVM8Kw/export?format=txt"
def fetch_system_prompt_from_gdoc(avatar_id, system_prompt_url):
    print(" Updating system prompt for avatar: " + avatar_id + "...")
    prompt_dir = "avatars_context/" + avatar_id + "/prompt/"
    os.makedirs(prompt_dir, exist_ok=True)
    # url = "https://docs.google.com/document/d/1NYOOy8KkaLDBwvHvEVg1hVDY5yvHeLACUpCEkJVM8Kw/export?format=txt"
    doc_id = system_prompt_url.split("/document/d/", 1)[1].split("/", 1)[0]
    system_prompt_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    response = requests.get(system_prompt_url)
    response.raise_for_status()
    prompt = response.text.strip()
    # prompt = prompt[:prompt.find('General Internal Impressions')]

    # base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(prompt_dir, "system_prompt.txt")
    with open(file_path, "w") as f:
        f.write(prompt)
    print(" Done.")


def convert_docx_to_txt_and_cleanup(folder_path):
    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith(".docx") or "." not in file:
                try:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    txt_filename = os.path.splitext(file)[0] + ".txt"
                    txt_path = os.path.join(root, txt_filename)
                    with open(txt_path, "w", encoding="utf-8") as f:
                        f.write(text)
                    os.remove(file_path)
                    print(f"✅ Converted and deleted: {file_path}")
                except Exception as e:
                    print(f"❌ Failed to convert {file_path}: {e}")


def fetch_youtube_transcript(url, languages=["de"]):
    print(f"🔗 Fetching: {url}")
    try:
        parsed = urlparse(url)

        # Handle standard and short YouTube URL formats
        if "youtube.com" in parsed.netloc:
            video_id = parse_qs(parsed.query).get("v", [None])[0]
        elif "youtu.be" in parsed.netloc:
            video_id = parsed.path.lstrip("/")
        else:
            raise ValueError("Unrecognized YouTube URL format.")

        if not video_id or len(video_id) != 11:
            raise ValueError("Invalid YouTube video ID.")

        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=languages)
        full_text = " ".join([entry["text"] for entry in transcript])
        return LlamaDocument(text=full_text, metadata={"source": url})

    except Exception as e:
        print(f"❌ Failed to fetch {url}: {e}")
        return None


def sanitize_filename(url):
    domain = urlparse(url).netloc
    hashed = hashlib.md5(url.encode()).hexdigest()[:8]
    return f"{domain.replace('.', '_')}_{hashed}.txt"


# def get_llm(mode='openai',model_name=None, system_prompt=None):
#     # base_dir = os.path.dirname(os.path.abspath(__file__))
#     # file_path = os.path.join(base_dir, 'system_prompt.txt')
#     # if system_prompt == None:
#     #     system_prompt = open(file_path, 'r').read()

#     # if model_name != None:
#     if mode == 'openai':

#         llm =  OpenAI(
#             # point at your custom endpoint:
#             api_key=API_KEY,             # e.g. 'sk-…'
#             base_url=API_BASE,           # "https://llm.hrz.uni-giessen.de/api/"
#             # api_type="open_ai",          # use the “open_ai” protocol
#             # api_version=None,            # leave None unless your server needs a version
#         )


#     # else:

#     #     llm = GWDGChatLLM(
#     #             model=model_name,
#     #             api_base=API_BASE,
#     #             api_key=API_KEY,
#     #             temperature=0.5,
#     #             system_prompt=system_prompt
#     #         )

#     # print('LLM details: ', llm.model_dump())

#     # Settings.llm = llm

#     return llm #, system_prompt


def create_session_log():
    os.makedirs(LOG_DIR, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    return open(os.path.join(LOG_DIR, f"session_{timestamp}.txt"), "w")


def ensure_punkt_tab():
    try:
        # this will raise LookupError if not found
        nltk.data.find("tokenizers/punkt_tab")
    except LookupError:
        print("punkt_tab not found—downloading…")
        nltk.download("punkt_tab")


CHUNK_SIZE, OVERLAP = 200, 30


# ------------------------------------------------------------------
# 0. normalisers / helpers
# ------------------------------------------------------------------
def normalise(txt: str) -> str:
    txt = re.sub(r"\s+", " ", txt.lower().strip())
    return unicodedata.normalize("NFKD", txt)


def tokenize(text: str, lang: str) -> list[str]:
    ensure_punkt_tab()
    lang_flag = "german" if lang == "de" else "english"
    return word_tokenize(text, language=lang_flag)


# ------------------------------------------------------------------
# 1. translation helper with cache
# ------------------------------------------------------------------
_trans_cache: dict[tuple[str, str], str] = {}  # (text, target_lang) → translated


def translate(text: str, target_lang: str) -> str:
    """
    Translate 'text' to target language ('de' or 'en') using deep_translator.
    Caches results to avoid hitting rate limits.
    """
    key = (text, target_lang)
    if key in _trans_cache:
        return _trans_cache[key]

    translated = GoogleTranslator(source="auto", target=target_lang).translate(text)
    _trans_cache[key] = translated
    return translated


def prepare_text_index(RAW_TEXT):
    if not RAW_TEXT or not RAW_TEXT.strip():
        print("➡ No text content to index, returning empty text index.")
        return None, []

    ensure_punkt_tab()
    sentences = sent_tokenize(normalise(RAW_TEXT), language="german")
    chunks, cur, wc = [], [], 0
    for sent in sentences:
        words = tokenize(sent, "de")  # punkt model is DE but works for EN too
        cur.extend(words)
        wc += len(words)
        if wc >= CHUNK_SIZE:
            chunks.append(" ".join(cur))
            cur, wc = cur[-OVERLAP:], len(cur)
    if cur:
        chunks.append(" ".join(cur))

    print(f"➡  Raw chunks: {len(chunks)}")
    # ------------------------------------------------------------------
    # 2. Build your BM-25 index (assumes you already have 'chunks')
    # ------------------------------------------------------------------
    #   chunks = ["..."]    # list of cleaned 200-word blocks
    token_lists = [tokenize(c, detect(c)) for c in chunks]
    bm25 = BM25Okapi(token_lists)
    return bm25, chunks


# ------------------------------------------------------------------
# 3. Dual-language search
# ------------------------------------------------------------------
def search_text_index(bm25, chunks, en_keywords, de_keywords, k_each: int = 5):
    if bm25 is None:
        return []

    # keyword_list = query.split(', ')
    query = " ".join(en_keywords)  # keyword_list[:3])
    trans_q = " ".join(de_keywords)  # keyword_list[3:])

    lang_orig = "en"  # "de" if detect(query) == "de" else "en"
    lang_trans = "de"  # "en" if lang_orig == "de" else "de"

    # --- original language pass -----------------------------------
    q_tokens_o = tokenize(normalise(query), lang_orig)
    # print('Query recieved by Text Index searcher: ', ', '.join(en_keywords + de_keywords))
    # print('Keywords group 1: ', query)
    # print('Keywords group 2: ', trans_q)

    # print('Group 1 tokens to search with BM25: ', q_tokens_o)
    scores_o = bm25.get_scores(q_tokens_o)
    top_o = scores_o.argsort()[-k_each:][::-1]

    # --- translated pass ------------------------------------------
    # trans_q    = translate(query, lang_trans)
    # print('Translated query: ',trans_q)
    q_tokens_t = tokenize(normalise(trans_q), lang_trans)
    # print('Group 2 tokens to search with BM25: ', q_tokens_t)
    scores_t = bm25.get_scores(q_tokens_t)
    top_t = scores_t.argsort()[-k_each:][::-1]

    # --- merge, preferring best score if overlap ------------------
    seen, results = {}, []
    for idx in top_o:
        seen[idx] = ("orig", float(scores_o[idx]))
    for idx in top_t:
        if idx in seen:
            # keep the better score
            seen[idx] = ("orig+trans", max(seen[idx][1], float(scores_t[idx])))
        else:
            seen[idx] = ("trans", float(scores_t[idx]))

    # sort by score descending and trim to k_each*2
    merged = sorted(seen.items(), key=lambda kv: kv[1][1], reverse=True)[: 2 * k_each]
    for idx, (tag, score) in merged[:6]:
        # results.append((tag, score, chunks[idx]))
        results.append(chunks[idx])
    return results


def translate_keywords_batch(keywords, source_lang="auto", target_lang="de"):
    if target_lang == "de":
        prefix = "Word: "
    elif target_lang == "en":
        prefix = "Wort: "

    translator = GoogleTranslator(source=source_lang, target=target_lang)
    joined = ". ".join([prefix + w for w in keywords])
    out = translator.translate(joined)

    # Extract translations after 'Wort:' or similar
    parts = [
        p.strip().replace(".", "")
        for p in out.replace("Wort:", "Word:").split("Word:")
        if p.strip()
    ]
    return parts


# Lazy-loading NLP models
nlp_models = {}
nlp_fallback = None

def download_spacy_model(model_name):
    """Download a spacy model if not available."""
    import subprocess
    import sys

    print(f"Downloading spacy model: {model_name}...")

    try:
        # Try using spacy's download command first
        subprocess.run(
            [sys.executable, "-m", "spacy", "download", model_name],
            check=True,
            capture_output=True
        )
        print(f"Successfully downloaded {model_name}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"Failed to download {model_name} using spacy download: {e}")

        # Fallback: try pip install with direct URL
        # Mapping of model names to their download URLs (for small models)
        model_urls = {
            'en_core_web_sm': 'https://github.com/explosion/spacy-models/releases/download/en_core_web_sm-3.7.0/en_core_web_sm-3.7.0-py3-none-any.whl',
            'de_core_news_sm': 'https://github.com/explosion/spacy-models/releases/download/de_core_news_sm-3.7.0/de_core_news_sm-3.7.0-py3-none-any.whl',
            'fr_core_news_sm': 'https://github.com/explosion/spacy-models/releases/download/fr_core_news_sm-3.7.0/fr_core_news_sm-3.7.0-py3-none-any.whl',
            'es_core_news_sm': 'https://github.com/explosion/spacy-models/releases/download/es_core_news_sm-3.7.0/es_core_news_sm-3.7.0-py3-none-any.whl',
            'pt_core_news_sm': 'https://github.com/explosion/spacy-models/releases/download/pt_core_news_sm-3.7.0/pt_core_news_sm-3.7.0-py3-none-any.whl',
        }

        if model_name in model_urls:
            try:
                subprocess.run(
                    [sys.executable, "-m", "pip", "install", model_urls[model_name]],
                    check=True,
                    capture_output=True
                )
                print(f"Successfully downloaded {model_name} via pip")
                return True
            except subprocess.CalledProcessError:
                print(f"Failed to download {model_name} via pip")
                return False

        return False

def get_nlp_model(lang_code):
    """Lazy-load NLP models on demand, downloading if necessary."""
    global nlp_models, nlp_fallback

    # Return cached model if available
    if lang_code in nlp_models:
        return nlp_models[lang_code]

    # Map language codes to spacy model names
    model_map = {
        'en': 'en_core_web_sm',
        'de': 'de_core_news_sm',
        'fr': 'fr_core_news_sm',
        'es': 'es_core_news_sm',
        'pt': 'pt_core_news_sm',
    }

    if lang_code in model_map:
        model_name = model_map[lang_code]

        # Try to load the model
        try:
            print(f"Loading spacy model for {lang_code} ({model_name})...")
            nlp_models[lang_code] = spacy.load(model_name)
            return nlp_models[lang_code]
        except OSError:
            print(f"Spacy model {model_name} not found, attempting to download...")
            if download_spacy_model(model_name):
                # Try loading again after download
                try:
                    nlp_models[lang_code] = spacy.load(model_name)
                    print(f"Successfully loaded {model_name} after download")
                    return nlp_models[lang_code]
                except OSError:
                    print(f"Still cannot load {model_name} after download, using fallback")

    # Use multilingual fallback
    if nlp_fallback is None:
        print("Loading multilingual spacy model (xx_ent_wiki_sm)...")
        try:
            nlp_fallback = spacy.load("xx_ent_wiki_sm")
        except OSError:
            print("Multilingual model not available, using English as fallback")
            # Ensure English model is available (download if needed)
            if 'en' not in nlp_models:
                get_nlp_model('en')
            nlp_fallback = nlp_models.get('en', spacy.load('en_core_web_sm'))

    return nlp_fallback


# Preload commonly used models
print("Preloading EN and DE spacy models...")
nlp_en = get_nlp_model('en')
nlp_de = get_nlp_model('de')


def extract_keywords(text):
    # Detect dominant language
    lang, confidence = langid.classify(text)
    print("Processing: ", text)
    print(f"Detected language: {lang} (confidence={confidence:.2f}). Input: {text}")

    # Decide model and translation target
    if lang.startswith("de"):
        nlp = nlp_de
        target_lang = "en"
    elif lang.startswith("en"):
        nlp = nlp_en
        target_lang = "de"
    else:
        # Multilingual or low-confidence input
        text = GoogleTranslator(source="auto", target="en").translate(text)
        nlp = nlp_en
        target_lang = "de"

    doc = nlp(text)
    print("extract_keyword doc: ", doc)
    # Extract nouns and proper nouns as keywords
    keywords = [
        token.text
        for token in doc
        if token.pos_ in ("NOUN", "PROPN", "X")
        and (not token.is_stop or token.pos_ == "PROPN")
    ]  # if (token.pos_ in ("NOUN", "PROPN", "X") or (token.text[0].isupper() and token.i != 0)) and not token.is_stop]
    keywords = list(set(keywords))  # remove duplicates
    if len(keywords) == 0:
        keywords = text.split(" ")
    print("Keywords: ", keywords)
    # Handle empty keywords ***

    translated_keywords = translate_keywords_batch(keywords, target_lang=target_lang)
    print("Translated keywords: ", translated_keywords)

    if target_lang == "en":
        en_keywords = translated_keywords
        de_keywords = keywords
    else:
        en_keywords = keywords
        de_keywords = translated_keywords

    return en_keywords, de_keywords


def extract_keywords_multilingual(text, rag_languages):
    """
    Extract keywords from user input and translate to RAG languages.

    Strategy:
    - Detect input language
    - If input is a RAG language: extract in that language, translate only to OTHER RAG languages
    - If not: use multilingual extraction, translate to ALL RAG languages
    """
    # Detect input language
    input_lang, confidence = langid.classify(text)
    print(f"Detected input language: {input_lang} (confidence={confidence})")

    result = {}

    # Check if input language is in RAG languages (or close enough)
    input_in_rag = any(
        input_lang == lang or input_lang.startswith(lang[:2])
        for lang in rag_languages
    )

    if input_in_rag:
        # Extract keywords in input language using appropriate model
        nlp = get_nlp_model(input_lang)

        doc = nlp(text)
        keywords = [
            token.text for token in doc
            if token.pos_ in ("NOUN", "PROPN", "X")
            and (not token.is_stop or token.pos_ == "PROPN")
        ]
        keywords = list(set(keywords))[:5]
        if len(keywords) == 0:
            keywords = text.split()[:5]

        result[input_lang] = keywords

        # Translate only to OTHER RAG languages (not input lang)
        for target_lang in rag_languages:
            if target_lang != input_lang:
                try:
                    translated = translate_keywords_batch(keywords, target_lang=target_lang)
                    result[target_lang] = translated
                except Exception as e:
                    print(f"Translation to {target_lang} failed: {e}")
                    result[target_lang] = keywords
    else:
        # Input language not in RAG - use multilingual extraction
        nlp = get_nlp_model(input_lang)  # Will use fallback if not available

        doc = nlp(text)
        keywords = [
            token.text for token in doc
            if token.pos_ in ("NOUN", "PROPN", "X")
            and (not token.is_stop or token.pos_ == "PROPN")
        ]
        keywords = list(set(keywords))[:5]
        if len(keywords) == 0:
            keywords = text.split()[:5]

        # Translate to ALL RAG languages
        for target_lang in rag_languages:
            try:
                translated = translate_keywords_batch(keywords, target_lang=target_lang)
                result[target_lang] = translated
            except Exception as e:
                print(f"Translation to {target_lang} failed: {e}")
                result[target_lang] = keywords

    return result


def search_text_index_multilingual(bm25, chunks, keywords_by_lang, k_each=3):
    """Search BM25 index with multiple language keyword sets."""
    if bm25 is None:
        return []

    all_scores = {}

    for lang, keywords in keywords_by_lang.items():
        query = " ".join(keywords)
        # Use 'de' tokenization as default (works for most European languages)
        lang_code = "de" if lang == "de" else "en"

        tokens = tokenize(normalise(query), lang_code)
        scores = bm25.get_scores(tokens)
        top = scores.argsort()[-k_each:][::-1]

        for idx in top:
            if idx in all_scores:
                all_scores[idx] = max(all_scores[idx], float(scores[idx]))
            else:
                all_scores[idx] = float(scores[idx])

    # Sort by score and return top results
    merged = sorted(all_scores.items(), key=lambda kv: kv[1], reverse=True)[:k_each*2]
    return [chunks[idx] for idx, score in merged]


# Sensor related


# 1) Fetch & normalize your ThingSpeak data
THINGSPEAK_URL = "https://api.thingspeak.com/channels/2974588/feeds.json?results=100"


# 2) Wrap it in a callable that runs PandasQueryEngine on demand
class SensorsTool:
    name = "sensors"
    description = "Access live sensor data and perform analysis on environmental readings."

    def __init__(self, llm, sensor_url=None, sensor_description=None):
        # store whichever LLM you pass in (e.g. get_llm("mistral-large-instruct"))
        self.llm = llm
        self.cache_ttl = 1800
        self._cached_df = None
        self._engine = None

        # Store custom URL and description if provided
        self.sensor_url = sensor_url
        self.sensor_description = sensor_description
        self.GENERAL_PANDAS_INSTRUCTIONS = """
            General guidance for using the DataFrame `df`:

            CRITICAL: Output ONLY raw Python code. Do NOT use markdown code blocks (no ```python, no ```).
            The code execution system cannot parse markdown syntax.

            1. Always ensure consistent datetime handling:
               df['created_at'] = pd.to_datetime(df['created_at'])
               df['created_at'] = df['created_at'].dt.tz_localize(None)
               df = df.set_index('created_at')

            2. Never assume exact timestamp equality; use nearest lookup:
               idx = df.index.get_indexer([target_time], method='nearest')[0]

            3. Handle NaN or missing values gracefully:
               Use df.fillna(method='ffill') or df.interpolate() as needed.

            4. Guard against missing columns:
               Always check `"Temp (°C)" in df.columns` before accessing.

            5. For aggregation or change-over-time questions:
               Choose the largest of [seconds, minutes, hours, days, weeks, months, years]
               smaller than the requested interval. Use resampling accordingly.

            6. If any error occurs, adapt your next code so that it handles the failure case robustly.
            """

    def _fetch_sensor_data_df(self) -> pd.DataFrame:
        print(f"Fetching sensor data from {self.sensor_url}...")
        headers = {
            "Accept": "application/json",
            "User-Agent": "Lahn-Avatar/1.0"
        }
        resp = requests.get(self.sensor_url, headers=headers)
        resp.raise_for_status()
        data = resp.json()

        # Check if this is ThingSpeak format or other format
        if "channel" in data:
            # ThingSpeak format: has "channel" metadata with field definitions
            channel_meta = data["channel"]
            field_map = {f"field{i}": channel_meta[f"field{i}"] for i in range(1, 7) if f"field{i}" in channel_meta}

            # load feeds into DataFrame
            df = pd.json_normalize(data["feeds"])

            # rename columns to friendly names from metadata
            df = df.rename(columns=field_map)

            # parse timestamp & convert all sensor readings to numeric
            df["created_at"] = pd.to_datetime(df["created_at"])
            for col in field_map.values():
                df[col] = pd.to_numeric(df[col], errors="coerce")
        else:
            # Generic format: use feeds as-is
            # (e.g., AIRBOX format with s_d0, s_d1, s_t0, etc.)
            df = pd.json_normalize(data["feeds"])

            # Parse timestamp if available
            if "timestamp" in df.columns:
                df["created_at"] = pd.to_datetime(df["timestamp"])
            elif "created_at" in df.columns:
                df["created_at"] = pd.to_datetime(df["created_at"])
            else:
                # No timestamp column, use index
                df.reset_index(drop=True, inplace=True)

            # Convert numeric columns
            for col in df.columns:
                if col.startswith("s_") or col.startswith("c_"):
                    df[col] = pd.to_numeric(df[col], errors="coerce")

        return df

    def _get_df(self):
        now = time.time()
        if self._cached_df is None or (now - self._last_fetch) > self.cache_ttl:
            print("Fetching new sensor data...")
            self._cached_df = self._fetch_sensor_data_df()
            self._last_fetch = now
        return self._cached_df

    def _get_df_sample(self, n=10):
        """Return a representative text sample of the dataframe for the LLM."""
        import pandas as pd

        df = self._get_df()
        sample = df.head(n).to_string(index=False)
        info = f"DataFrame columns: {', '.join(df.columns)}\nSample rows:\n{sample}"
        return info

    def __call__(self, query: str) -> str:
        print("Calling Sensor Tool...")
        n_tries = 0
        df = self._get_df()
        sample_info = self._get_df_sample()

        if self._engine is None:
            self._engine = PandasQueryEngine(
                df=df, llm=self.llm, verbose=True, synthesize_response=False
            )
        else:
            self._engine.df = df

        query = (
            f"{query}\n\n"
            f"{self.GENERAL_PANDAS_INSTRUCTIONS}\n\n"
            f"Context: here is a small sample of the dataframe you will analyze.\n"
            f"Use this to infer datetime granularity, column names, and data structure.\n"
            f"{sample_info}\n"
            # "If you need to access a specific timestamp, use the nearest available one "
            # "rather than assuming exact equality. "
            "Sometimes values are null, so take that into account as well."
        )

        result = self._engine.query(query).response

        # If the response contains an embedded Pandas failure message, trigger repair
        while (isinstance(result, str)) and ("Error message:" in result):
            if n_tries < 3:
                n_tries += 1
                print(
                    "⚠️ Detected embedded error message in response — retrying with augmented query..."
                )
                query = f"{query}\n\nNote: the previous code failed with this error: {result}. Identify the reason for the error, and adapt your new approach to avoid that."
                print("Augmented query: ", query)
                result = self._engine.query(query).response
            else:
                print("Unable to analyze data after multiple tries.")
                return "Technical issue with sensor data analysis. Pls try again later."
        return result

    def query(self, query_str: str) -> str:
        """
        Alias so that QueryEngineTool can call .query(...)
        under the hood. Simply forwards to __call__.
        """
        return self(query_str)


def format_history_as_string(history):
    # print('To convert to string. Input: ', history)

    role_map = {"user": "User", "avatar": "Lahn"}

    result = "\n".join(
        f"{role_map.get(m['sender'], m['sender'])}: {m['text']}" for m in history
    )

    # print('Converted conversation history into string: ', result)

    return result


def convert_to_wav(input_path, output_path):
    command = [
        "ffmpeg",
        "-y",
        "-i",
        input_path,
        "-ar",
        "16000",
        "-ac",
        "1",
        output_path,
    ]
    subprocess.run(command, check=True)


def pcm_to_wav_bytes(pcm_bytes, sample_rate=24000, n_channels=1, sampwidth=2):
    buf = io.BytesIO()
    with wave.open(buf, "wb") as wf:
        wf.setnchannels(n_channels)
        wf.setsampwidth(sampwidth)  # 2 bytes for int16
        wf.setframerate(sample_rate)
        wf.writeframes(pcm_bytes)
    return buf.getvalue()


def transcribe_audio(file_path):
    from transformers import WhisperForConditionalGeneration, WhisperProcessor

    whisper_device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"🔄 Loading Whisper model on {whisper_device}...")
    whisper_processor = WhisperProcessor.from_pretrained("openai/whisper-small")
    whisper_model = WhisperForConditionalGeneration.from_pretrained(
        "openai/whisper-small"
    ).to(whisper_device)
    print("✅ Whisper model loaded.")

    temp_wav_path = file_path.rsplit(".", 1)[0] + "_converted.wav"
    convert_to_wav(file_path, temp_wav_path)

    speech, sr = torchaudio.load(temp_wav_path)
    input_features = whisper_processor(
        speech.squeeze(), sampling_rate=sr, return_tensors="pt"
    ).input_features.to(whisper_device)

    predicted_ids = whisper_model.generate(input_features)
    transcription = whisper_processor.batch_decode(
        predicted_ids, skip_special_tokens=True
    )[0]
    return transcription


# RAG related


def build_index(avatar_id, drive_folder_id):
    index_dir = "avatars_context/" + avatar_id + "/index"
    data_dir = "avatars_context/" + avatar_id + "/data"

    print("Contents of " + data_dir + " :")
    for root, dirs, files in os.walk(data_dir):
        for name in files:
            print(os.path.join(root, name))

    print("\nRefreshing from Google Drive...")
    download_drive_folder(drive_folder_id, data_dir)
    convert_docx_to_txt_and_cleanup(data_dir)

    print(f"\n--- Verifying contents of {data_dir} ---")
    for root, dirs, files in os.walk(data_dir):
        for name in files:
            print(f"Found file: {os.path.join(root, name)}")
        for name in dirs:
            print(f"Found dir: {os.path.join(root, name)}")
    print("--- Verification complete ---\n")

    print("\n\nCreating Context store from data sources...")

    documents = []
    if len(os.listdir(data_dir)) > 0:
        documents = SimpleDirectoryReader(data_dir, recursive=True).load_data()
        print(f"{len(documents)} documents loaded from {data_dir}")
        # for i, doc in enumerate(documents):
        #     print(f"\n--- Document {i+1} ---")
        #     print("File:", doc.metadata.get('file_path', 'Unknown'))
        #     print("Content preview:", doc.text[:300], "...\n")

    links_path = Path(data_dir) / "General_News/Online News (Links).txt"
    if links_path.exists():
        with open(links_path, "r") as f:
            urls = [line.strip() for line in f if line.strip()]

        web_reader = SimpleWebPageReader()
        for url in urls:
            try:
                print(f"🔗 Fetching: {url}")
                if "youtube.com" in url or "youtu.be" in url:
                    doc = fetch_youtube_transcript(url)
                else:
                    docs = web_reader.load_data([url])
                    full_text = "\n\n".join(doc.text for doc in docs)
                    doc = LlamaDocument(text=full_text, metadata={"source": url})

                if doc:
                    filename = sanitize_filename(url)
                    filepath = Path(data_dir) / "General_News/scraped_texts" / filename
                    filepath.parent.mkdir(parents=True, exist_ok=True)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(doc.text)
                    print(f"✅ Saved to {filepath}")
            except Exception as e:
                print(f"❌ Failed to fetch {url}: {e}")

    scraped_documents_path = Path(data_dir) / "General_News/scraped_texts"
    if (
        scraped_documents_path.exists()
        and len(os.listdir(str(scraped_documents_path))) > 0
    ):
        scraped_documents = SimpleDirectoryReader(
            str(scraped_documents_path)
        ).load_data()
        documents += scraped_documents

    experiences_folder_path = Path(data_dir) / "uploaded_experiences/text"
    if (
        experiences_folder_path.exists()
        and len(os.listdir(str(experiences_folder_path))) > 0
    ):
        new_uploads = SimpleDirectoryReader(
            str(experiences_folder_path), recursive=True
        ).load_data()
        documents += new_uploads

    vector_index = VectorStoreIndex.from_documents(documents)
    vector_index.storage_context.persist(persist_dir=index_dir)

    context = "\n".join([doc.text for doc in documents])
    text_index, chunks = prepare_text_index(context)

    pickle.dump(text_index, open(index_dir + "/text_index.pkl", "wb"))
    pickle.dump(chunks, open(index_dir + "/chunks.pkl", "wb"))

    print("Done")

    return vector_index, text_index, chunks


def build_or_load_index(avatar_id, drive_folder_id, refresh=False):
    index_dir = "avatars_context/" + avatar_id + "/index"
    data_dir = "avatars_context/" + avatar_id + "/data"

    index_ready = (
        os.path.exists(index_dir)
        # Vector index
        and os.path.exists(os.path.join(index_dir, "docstore.json"))
        and os.path.exists(os.path.join(index_dir, "index_store.json"))
        # Text index
        and os.path.exists(os.path.join(index_dir, "text_index.pkl"))
        and os.path.exists(os.path.join(index_dir, "chunks.pkl"))
    )

    if index_ready and not refresh:
        print("Loading index from storage...")
        storage_context = StorageContext.from_defaults(persist_dir=index_dir)
        vector_index = load_index_from_storage(storage_context)

        text_index = pickle.load(open(index_dir + "/text_index.pkl", "rb"))
        chunks = pickle.load(open(index_dir + "/chunks.pkl", "rb"))
        return vector_index, text_index, chunks

    # Index needs to be built and loaded
    vector_index, text_index, chunks = build_index(avatar_id, drive_folder_id)

    return vector_index, text_index, chunks


def prepare_query_engines(avatar_id, drive_folder_id, rag_languages=None, refresh=False):
    if drive_folder_id == None:
        return [None, None, None, None, rag_languages or ['en', 'de']]

    # Default to EN/DE if not specified
    if rag_languages is None:
        rag_languages = ['en', 'de']

    if refresh == True:
        vector_index, text_index, chunks = build_index(avatar_id, drive_folder_id)
    else:
        vector_index, text_index, chunks = build_or_load_index(
            avatar_id, drive_folder_id
        )

    # query_llm = get_llm('gwdg', "mistral-large-instruct", system_prompt= 'Provide an accurate response to the given query:')
    vector_index_query_engine = vector_index.as_retriever(
        similarity_top_k=5, verbose=True
    )
    # vector_index_query_engine = vector_index.as_query_engine(llm=vector_query_llm,similarity_top_k=10, verbose=True)
    text_index_query_engine = search_text_index

    return vector_index_query_engine, text_index_query_engine, text_index, chunks, rag_languages


# vector_index_query_engine, text_index_query_engine, text_index, chunks = prepare_query_engines()


def fetch_text_index_query(text_query_llm, conversation):
    print("Fetching context from text index...")
    query_prompt = "Here is the conversation: " + format_history_as_string(
        conversation
    )  # + '\nUser: '+prompt #response[:response.find('")')]
    # print('Query prompt: ', query_prompt)

    query = str(text_query_llm.complete(query_prompt))
    # print('Crafted Query: ', query)
    return query


def generate_context_aware_keywords_for_multilingual_text_index_search(text_query_llm, conversation, rag_languages):
    """
    Analyze the full conversation context and generate keywords for multilingual text-index (BM25) search.

    This function looks beyond just the last message - it examines the entire
    conversation to understand what context is needed for text-index search.
    This is especially important for follow-up questions where the search needs
    to maintain context from previous exchanges.

    NOTE: These keywords are specifically for the BM25 text-index search, NOT
    for vector search. Vector search uses all keywords combined.

    Examples:
        - First question "What about water quality?" -> ['water', 'quality']
        - Follow-up "And temperature?" -> ['temperature', 'water body', 'recent']
          (understands this is still about the same water body)

    Args:
        text_query_llm: LLM with system_prompt for query generation
        conversation: Full chat history for context analysis
        rag_languages: List of language codes (e.g., ['en', 'de', 'pt'])

    Returns:
        Dictionary mapping language codes to keyword lists
        e.g., {'en': ['water', 'quality'], 'pt': ['água', 'qualidade']}
    """
    print(f"Generating context-aware keywords for RAG languages: {rag_languages}")

    # Build language-specific prompt
    languages_str = ", ".join(rag_languages)
    num_keywords = 3  # Keywords per language

    # Create dynamic system prompt based on RAG languages
    system_prompt = f"""Context is needed to address the most recent message in the conversation.
Look through the given conversation and determine what context is needed.

If specific context is needed, return {num_keywords} relevant keywords for each of these languages: {languages_str}.
If no specific context is needed (e.g., simple greetings), return general keywords about the Lahn/river.

Also determine whether the user's last message requires live/current information from the internet
(e.g. recent news, recent research publications, current events not in a static knowledge base).
Do NOT flag as web search: questions answerable from the knowledge base, sensor readings, or general conversation.

Return your response in exactly TWO lines:
Line 1: keyword1_lang1, keyword2_lang1, keyword3_lang1 | keyword1_lang2, keyword2_lang2, keyword3_lang2 | ...
Line 2: WEB_SEARCH: <specific search query>   (if live internet data is needed)
         WEB_SEARCH: NONE                       (otherwise)

Where languages on line 1 are in this order: {languages_str}

For example, if languages are ['en', 'de']:
"water, quality, measurement | Wasser, Qualität, Messung
WEB_SEARCH: NONE"

Or if recent news is needed:
"Lahn, news, research | Lahn, Nachrichten, Forschung
WEB_SEARCH: latest Lahn river research 2026"

Reply only in this two-line format, nothing else."""

    # Update the LLM's system prompt temporarily
    original_prompt = text_query_llm.system_prompt
    text_query_llm.system_prompt = system_prompt

    try:
        query_prompt = "Here is the conversation: " + format_history_as_string(conversation)
        response = str(text_query_llm.complete(query_prompt)).strip()

        print(f"LLM generated keywords + web intent: {response}")

        lines = response.splitlines()
        keywords_line = lines[0].strip() if lines else ""
        web_line = lines[1].strip() if len(lines) > 1 else ""

        # Parse RAG keywords
        keywords_by_lang = {}
        parts = keywords_line.split("|")
        for i, part in enumerate(parts):
            if i < len(rag_languages):
                lang = rag_languages[i]
                keywords = [k.strip() for k in part.split(",") if k.strip()]
                keywords_by_lang[lang] = keywords[:num_keywords]

        for lang in rag_languages:
            if lang not in keywords_by_lang:
                keywords_by_lang[lang] = []

        # Parse web search intent
        web_search_query = None
        if web_line.upper().startswith("WEB_SEARCH:"):
            q = web_line[len("WEB_SEARCH:"):].strip()
            if q.upper() != "NONE" and q:
                web_search_query = q

        print(f"Parsed keywords by language: {keywords_by_lang}")
        print(f"Web search intent: {web_search_query!r}")
        return keywords_by_lang, web_search_query

    finally:
        # Restore original system prompt
        text_query_llm.system_prompt = original_prompt


def detect_web_search_intent(text_query_llm, conversation):
    """
    Lightweight call to determine whether the user's last message requires a live web search.
    Used for avatars that have no RAG index (so no keyword-generation call is made).
    Returns a search query string, or None.
    """
    system_prompt = """Does the user's last message require live/current information from the internet?
(e.g. recent news, recent research, current events not in a static knowledge base)
Do NOT flag: questions about the Lahn answerable from a knowledge base, current sensor readings, or general conversation.

Reply with exactly one line:
WEB_SEARCH: <specific search query>   — if live internet data is needed
WEB_SEARCH: NONE                       — otherwise"""

    original_prompt = text_query_llm.system_prompt
    text_query_llm.system_prompt = system_prompt
    try:
        response = str(text_query_llm.complete(
            "Here is the conversation: " + format_history_as_string(conversation)
        )).strip()
        print(f"Web search intent (no-RAG path): {response!r}")
        if response.upper().startswith("WEB_SEARCH:"):
            q = response[len("WEB_SEARCH:"):].strip()
            if q.upper() != "NONE" and q:
                return q
        return None
    finally:
        text_query_llm.system_prompt = original_prompt


def fetch_vector_index_context(vector_index_query_engine, query):
    print("Fetching context from vector index...")
    v_response = vector_index_query_engine.retrieve(query)
    response = "\n\n".join(r.node.text for r in v_response)
    # response =  vector_index_query_engine.query(query)
    # print('\n\nContext from vector index: ', response)

    response = "\n\n".join(
        "".join(
            ch for ch in r.node.text if 32 <= ord(ch) <= 126 or ch in "\n\r\t"
        )  # Why tf is there binary data in the results?
        for r in v_response
    )

    print("Done fetching context from vector index...")
    return response


def RAG(avatar_rag_tools, query=None, translated=False, keywords_by_lang=None):
    """
    Retrieve and augment context using RAG for the avatar's knowledge base.

    Args:
        avatar_rag_tools: Tuple of (vector_engine, text_search_fn, text_index, chunks, rag_languages)
        query: User query string (used only if keywords_by_lang is None)
        translated: Legacy parameter (deprecated)
        keywords_by_lang: Pre-generated context-aware keywords by language

    The keywords_by_lang parameter should be provided by generate_context_aware_keywords_for_multilingual_text_index_search
    for context-aware query generation that considers the full conversation.
    """
    # Get RAG languages from avatar tools (5th element)
    rag_languages = avatar_rag_tools[4] if len(avatar_rag_tools) > 4 else ['en', 'de']

    # Determine keywords based on what was provided
    if keywords_by_lang is not None:
        # Use pre-generated context-aware keywords (preferred path)
        # print(f"Using pre-generated context-aware keywords: {keywords_by_lang}")
        pass
    elif query is not None:
        # Extract keywords from the query directly (fallback)
        # print(f"Extracting keywords from query: {query}")
        keywords_by_lang = extract_keywords_multilingual(query, rag_languages)
    elif translated:
        # Legacy path for backward compatibility
        print(f"DEBUG: RAG function - translated=True, query='{query}'")
        keywords_by_lang = {}
        parts = query.split("|")
        if len(parts) >= 2:
            keywords_by_lang['en'] = [w.strip() for w in parts[0].split(",")]
            keywords_by_lang['de'] = [w.strip() for w in parts[1].split(",")]
    else:
        # No query or keywords provided
        print("WARNING: No query or keywords provided to RAG")
        keywords_by_lang = {lang: [] for lang in rag_languages}

    print(f"Keywords by language: {keywords_by_lang}")

    # Combine all keywords for vector search
    all_keywords = []
    for lang_keywords in keywords_by_lang.values():
        all_keywords.extend(lang_keywords)

    with ThreadPoolExecutor(max_workers=2) as executor:
        thread_0 = executor.submit(
            fetch_vector_index_context,
            avatar_rag_tools[0],
            ", ".join(all_keywords),
        )  # query)
        thread_1 = executor.submit(
            search_text_index_multilingual,
            avatar_rag_tools[2],  # BM25 object (text_index)
            avatar_rag_tools[3],  # chunks
            keywords_by_lang,
        )

        wait([thread_0, thread_1])

    # context_from_text_index = thread_0
    context_from_vector_index = thread_0.result()  # .response
    context_from_text_index = thread_1.result()
    context_from_text_index = "\n".join(context_from_text_index)

    total_context = (
        "\nContext from text-based retrieval: \n"
        + context_from_text_index
        + "\n------------\nContext from vector-based retrieval: \n"
        + context_from_vector_index
    )

    total_context = (
        "Here is relevant information (Sometimes the text-retrieval has relevant information that the vector-retrieval doesn't, or vice versa. Look through each comprehensively, to extract the most relevant information you need. Even if the Vector-retrieval says there's no information available, still scrutinize the Text-retrieval results to fetch relevant info. Also make sure to reply in the same language the user messaged you in -not necessarily the language in which this context is expressed. If the user messaged you in English, reply in English as well, even if this context is in German. If the user messaged you in German, respond in German, if it was Portuguese, respond in Portuguese, etc."
        + total_context
    )

    printable = " | ".join(
        line.strip() for line in total_context.splitlines() if line.strip()
    )
    # print("-----\n\nRAG result: ", printable)
    return printable


# ---------- Web Search (Brave API) ----------
def web_search(query: str, count: int = 5) -> str:
    """
    Search the web using Brave Search API.
    Brave automatically detects query language and returns matching results.

    Args:
        query: The search query
        count: Number of results to return (default 5, max 50)

    Returns:
        Formatted search results string for LLM consumption
    """
    api_key = os.getenv("BRAVE_SEARCH_API_KEY")
    if not api_key:
        return "Web search unavailable: API key not configured."

    url = "https://api.search.brave.com/res/v1/web/search"

    headers = {
        "Accept": "application/json",
        "Accept-Encoding": "gzip",
        "X-Subscription-Token": api_key
    }

    params = {
        "q": query,
        "count": count
    }

    try:
        resp = requests.get(url, headers=headers, params=params, timeout=10)
        resp.raise_for_status()

        data = resp.json()
        results = data.get("web", {}).get("results", [])

        if not results:
            return f"No results found for query: {query}"

        # Format results for LLM
        formatted = []
        for r in results:
            title = r.get("title", "No title")
            desc = r.get("description", "No description")
            url_link = r.get("url", "")
            formatted.append(f"- **{title}**\n  {desc}\n  Source: {url_link}")

        return "\n\n".join(formatted)

    except requests.exceptions.RequestException as e:
        return f"Web search failed: {str(e)}"
    except Exception as e:
        return f"Web search error: {str(e)}"
